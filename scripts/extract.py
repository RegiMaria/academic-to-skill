#!/usr/bin/env python3
"""
academic-to-skill — Extrator de texto para documentos acadêmicos e técnicos.

Este script faz UMA coisa: extrai o texto bruto do documento e salva em arquivo.
O Claude Code (SKILL.md) cuida de todo o resto — estrutura, seções, arquivos
acadêmicos extras, detecção de artigo científico, etc.

Ordem de tentativa para PDF:
  1. pdftotext (poppler-utils) — melhor qualidade, preserva layout
     Instalar: sudo apt install poppler-utils
  2. PyPDF2 — biblioteca Python, sem dependências externas
     Instalar: pip3 install PyPDF2
  3. pdfminer.six — fallback mais completo
     Instalar: pip3 install pdfminer.six

Ordem de tentativa para EPUB:
  1. ebooklib + BeautifulSoup4 — melhor qualidade
     Instalar: pip3 install ebooklib beautifulsoup4
  2. zipfile + html.parser — fallback usando apenas a biblioteca padrão

Outros formatos suportados:
  - Texto simples / Markdown / reStructuredText / AsciiDoc — leitura direta
  - HTML — BeautifulSoup4 se disponível, senão html.parser da stdlib
  - DOCX — python-docx se disponível, senão fallback ZIP/XML da stdlib
  - RTF — striprtf se disponível, senão fallback com regex
  - MOBI/AZW/AZW3 — Calibre ebook-convert se instalado

Arquivos gerados:
  <tempdir>/book_skill_work/full_text.txt  — texto completo extraído
  <tempdir>/book_skill_work/metadata.json  — estatísticas e metadados

Variável de ambiente:
  BOOK_SKILL_WORKDIR — sobrescreve o diretório de saída padrão
"""

import html
import html.parser
import importlib.util
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path

# ─── Diretórios de saída ────────────────────────────────────────────────────
# Por padrão usa a pasta temporária do sistema.
# Pode ser sobrescrito pela variável de ambiente BOOK_SKILL_WORKDIR.
OUTPUT_DIR = Path(
    os.environ.get(
        "BOOK_SKILL_WORKDIR",
        str(Path(tempfile.gettempdir()) / "book_skill_work"),
    )
)
OUTPUT_TEXT = OUTPUT_DIR / "full_text.txt"   # Texto extraído do documento
OUTPUT_META = OUTPUT_DIR / "metadata.json"   # Estatísticas e metadados

# ─── Constantes ─────────────────────────────────────────────────────────────
# Estimativa de tokens: ~0,75 palavras por token (aproximação para inglês/português)
WORDS_PER_TOKEN = 0.75

# Extensões suportadas por categoria
TEXT_EXTENSIONS = {".txt", ".text", ".md", ".markdown", ".rst", ".adoc", ".asciidoc"}
HTML_EXTENSIONS = {".html", ".htm", ".xhtml"}
CALIBRE_EBOOK_EXTENSIONS = {".mobi", ".azw", ".azw3"}
SUPPORTED_EXTENSIONS = {
    ".pdf", ".epub", ".docx", ".rtf",
    *TEXT_EXTENSIONS,
    *HTML_EXTENSIONS,
    *CALIBRE_EBOOK_EXTENSIONS,
}

# Mapeamento: módulo Python → pacote pip para sugestões de instalação
PYTHON_DEPENDENCIES = {
    "docling": "docling",
    "PyPDF2": "PyPDF2",
    "pdfminer": "pdfminer.six",
    "ebooklib": "ebooklib",
    "bs4": "beautifulsoup4",
    "docx": "python-docx",
    "striprtf": "striprtf",
}


# ─── Funções utilitárias ─────────────────────────────────────────────────────

def estimate_tokens(text: str) -> int:
    """Estima o número de tokens usando ~0,75 palavras por token."""
    return int(len(text.split()) / WORDS_PER_TOKEN)


def supported_formats_message() -> str:
    """Lista todos os formatos suportados."""
    return ", ".join(sorted(SUPPORTED_EXTENSIONS))


def python_module_available(module_name: str) -> bool:
    """Verifica se um módulo Python está instalado sem importá-lo."""
    return importlib.util.find_spec(module_name) is not None


def missing_python_packages(module_names: list[str]) -> list[str]:
    """Retorna pacotes pip que precisam ser instalados."""
    missing = []
    for module_name in module_names:
        if not python_module_available(module_name):
            missing.append(PYTHON_DEPENDENCIES[module_name])
    return missing


def install_python_packages(packages: list[str]) -> bool:
    """Instala pacotes pip no ambiente Python atual."""
    if not packages:
        return True
    print(f"Installing missing Python package(s): {', '.join(packages)}")
    try:
        result = subprocess.run(
            [sys.executable, "-m", "pip", "install", *packages],
            text=True,
            timeout=600,
        )
    except Exception as exc:
        print(f"Package installation failed: {exc}", file=sys.stderr)
        return False
    # Limpa o cache para que os pacotes recém-instalados sejam encontrados
    importlib.invalidate_caches()
    return result.returncode == 0


def normalize_install_mode(argv: list[str]) -> str:
    """Determina o modo de instalação de dependências.
    
    Prioridade: flag na linha de comando > variável de ambiente > padrão "ask"
    Valores válidos: "yes" (instala sem perguntar), "no" (usa fallback), "ask"
    """
    mode = os.environ.get("BOOK_SKILL_INSTALL_MISSING", "ask").lower()
    if "--no-install-missing" in argv:
        return "no"
    if "--install-missing" in argv:
        idx = argv.index("--install-missing")
        if idx + 1 < len(argv) and not argv[idx + 1].startswith("--"):
            mode = argv[idx + 1].lower()
        else:
            mode = "yes"
    if mode in {"1", "true", "y", "yes", "install"}:
        return "yes"
    if mode in {"0", "false", "n", "no", "fallback", "skip"}:
        return "no"
    return "ask"


def offer_dependency_install(
    *,
    feature: str,
    module_names: list[str],
    fallback: str | None,
    install_mode: str,
) -> None:
    """Oferece instalar dependências opcionais quando estão faltando.
    
    - modo "yes": instala automaticamente sem perguntar
    - modo "ask" com terminal interativo: pergunta ao usuário
    - modo "no" ou sem terminal: usa o fallback silenciosamente
    """
    packages = missing_python_packages(module_names)
    if not packages:
        return  # Todos os pacotes já estão instalados

    message = f"{feature} uses {', '.join(packages)} if installed"
    if fallback:
        message += f", otherwise {fallback}"
    message += "."
    print(message)

    should_install = False
    if install_mode == "yes":
        should_install = True
    elif install_mode == "ask" and sys.stdin.isatty():
        # Só pergunta se estivermos em um terminal interativo
        answer = input("Missing package(s) detected. Do you want to install? y=install, n=fallback: ").strip().lower()
        should_install = answer in {"y", "yes", "install"}
    else:
        if fallback:
            print("Non-interactive mode or install disabled; using fallback.")
        else:
            print("Non-interactive mode or install disabled; installation skipped.")

    if not should_install:
        if fallback:
            print(f"Using fallback: {fallback}.")
        return

    if install_python_packages(packages):
        still_missing = missing_python_packages(module_names)
        if not still_missing:
            print("Package installation complete.")
            return
        print(f"Package installation incomplete; still missing: {', '.join(still_missing)}", file=sys.stderr)
    else:
        print("Package installation failed.", file=sys.stderr)

    if fallback:
        print(f"Using fallback: {fallback}.")


def prepare_dependencies(ext: str, extraction_mode: str, install_mode: str) -> None:
    """Verifica e oferece instalar dependências para o formato do arquivo."""
    # PDF técnico: Docling preserva tabelas e fórmulas como markdown
    if ext == ".pdf" and extraction_mode == "technical":
        offer_dependency_install(
            feature="Technical PDF extraction",
            module_names=["docling"],
            fallback="the PDF text fallback chain",
            install_mode=install_mode,
        )
    # PDF texto: tenta PyPDF2 e pdfminer se pdftotext não estiver no sistema
    if ext == ".pdf" and not shutil.which("pdftotext"):
        offer_dependency_install(
            feature="PDF text extraction",
            module_names=["PyPDF2", "pdfminer"],
            fallback="any installed Python PDF parser; extraction fails if none are available",
            install_mode=install_mode,
        )
    # EPUB: ebooklib + bs4 para melhor qualidade; zipfile como fallback
    if ext == ".epub":
        offer_dependency_install(
            feature="EPUB extraction",
            module_names=["ebooklib", "bs4"],
            fallback="a stdlib ZIP/HTML parser",
            install_mode=install_mode,
        )
    # HTML: beautifulsoup4 para melhor limpeza; html.parser como fallback
    if ext in HTML_EXTENSIONS:
        offer_dependency_install(
            feature="HTML extraction",
            module_names=["bs4"],
            fallback="a stdlib HTML parser",
            install_mode=install_mode,
        )
    # DOCX: python-docx para extração completa incluindo tabelas
    if ext == ".docx":
        offer_dependency_install(
            feature="DOCX extraction",
            module_names=["docx"],
            fallback="a stdlib ZIP/XML parser",
            install_mode=install_mode,
        )
    # RTF: striprtf para limpeza correta; regex básico como fallback
    if ext == ".rtf":
        offer_dependency_install(
            feature="RTF extraction",
            module_names=["striprtf"],
            fallback="a basic regex cleanup fallback",
            install_mode=install_mode,
        )


# ─── Leitores de arquivo ─────────────────────────────────────────────────────

def read_text_file(path: str) -> str | None:
    """Lê arquivo de texto tentando múltiplos encodings.
    
    Ordem: utf-8-sig (BOM), utf-8, cp1252 (Windows PT), latin-1 (fallback universal)
    """
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return Path(path).read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
        except Exception:
            return None
    return None


# ─── Extratores HTML ─────────────────────────────────────────────────────────

def extract_html_content(raw_html: str) -> str:
    """Extrai texto limpo de HTML — usa BeautifulSoup4 se disponível."""
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(raw_html, "html.parser")
        for element in soup(["script", "style", "head"]):
            element.decompose()
        return soup.get_text(separator="\n")
    except ImportError:
        # Fallback: parser HTML minimalista sem dependências externas
        parser = _HTMLTextExtractor()
        parser.feed(raw_html)
        return parser.get_text()


def extract_html_file(path: str) -> str | None:
    """Lê arquivo HTML e extrai texto limpo."""
    raw = read_text_file(path)
    if raw is None:
        return None
    return extract_html_content(raw)


# ─── Extratores DOCX ─────────────────────────────────────────────────────────

def extract_docx_with_python_docx(docx_path: str) -> str | None:
    """Extrai texto de DOCX usando python-docx — inclui parágrafos e tabelas."""
    try:
        import docx
        document = docx.Document(docx_path)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        # Extrai tabelas — importante para artigos com tabelas de resultados
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
        return "\n".join(parts)
    except ImportError:
        return None
    except Exception:
        return None


def extract_docx_with_zipfile(docx_path: str) -> str | None:
    """Extrai texto de DOCX sem dependências — lê XML diretamente do ZIP."""
    try:
        import xml.etree.ElementTree as ET
        with zipfile.ZipFile(docx_path) as zf:
            xml_bytes = zf.read("word/document.xml")
        root = ET.fromstring(xml_bytes)
        # Namespace padrão do formato OOXML
        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        parts: list[str] = []
        for paragraph in root.iter(f"{namespace}p"):
            texts = [node.text for node in paragraph.iter(f"{namespace}t") if node.text]
            if texts:
                parts.append("".join(texts))
        return "\n".join(parts) if parts else None
    except Exception:
        return None


def extract_docx(docx_path: str) -> tuple[str, str]:
    """Extrai texto de DOCX tentando métodos em ordem de qualidade."""
    print("Trying python-docx...", end=" ", flush=True)
    text = extract_docx_with_python_docx(docx_path)
    if text and text.strip():
        print("OK")
        return text, "python-docx"
    print("not available")
    print("Trying stdlib DOCX parser...", end=" ", flush=True)
    text = extract_docx_with_zipfile(docx_path)
    if text and text.strip():
        print("OK")
        return text, "zipfile-docx"
    print("FAILED")
    print(
        "\nERROR: Could not extract text from DOCX.\n"
        "Install python-docx for best results:\n"
        "  pip3 install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)


# ─── Extratores RTF ──────────────────────────────────────────────────────────

def strip_rtf_fallback(raw: str) -> str:
    """Remove formatação RTF com regex — fallback sem dependências."""
    raw = re.sub(r"\\'[0-9a-fA-F]{2}", " ", raw)
    raw = re.sub(r"\\par[d]?", "\n", raw)
    raw = re.sub(r"\\tab", "\t", raw)
    raw = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", raw)
    raw = raw.replace("{", "").replace("}", "")
    return html.unescape(raw)


def extract_rtf(rtf_path: str) -> tuple[str, str]:
    """Extrai texto de RTF — tenta striprtf, depois regex como fallback."""
    raw = read_text_file(rtf_path)
    if raw is None:
        print("ERROR: Could not read RTF file", file=sys.stderr)
        sys.exit(1)
    try:
        from striprtf.striprtf import rtf_to_text
        text = rtf_to_text(raw)
        if text.strip():
            return text, "striprtf"
    except ImportError:
        pass
    except Exception:
        pass
    return strip_rtf_fallback(raw), "rtf-regex"


# ─── Extrator Calibre (MOBI/AZW) ─────────────────────────────────────────────

def extract_with_ebook_convert(input_path: str) -> str | None:
    """Converte ebook com Calibre — suporta MOBI, AZW, AZW3.
    
    Calibre é instalado como aplicação, não via pip.
    Download: https://calibre-ebook.com/download
    """
    if not shutil.which("ebook-convert"):
        return None
    output_path = OUTPUT_DIR / "ebook-convert-output.txt"
    try:
        result = subprocess.run(
            ["ebook-convert", input_path, str(output_path)],
            capture_output=True, text=True, timeout=300
        )
        if result.returncode == 0 and output_path.exists():
            text = output_path.read_text(encoding="utf-8", errors="replace")
            if text.strip():
                return text
    except Exception:
        pass
    return None


# ─── Extratores PDF ──────────────────────────────────────────────────────────

def extract_with_pdftotext(pdf_path: str) -> str | None:
    """Extrai texto com pdftotext — melhor qualidade para PDFs de texto.
    
    A flag -layout preserva o layout aproximado da página.
    Importante para PDFs de duas colunas (como artigos da Frontiers, IEEE).
    Instalar: sudo apt install poppler-utils
    """
    if not shutil.which("pdftotext"):
        return None
    try:
        result = subprocess.run(
            ["pdftotext", "-layout", pdf_path, "-"],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except Exception:
        pass
    return None


def extract_with_pypdf2(pdf_path: str) -> str | None:
    """Extrai texto com PyPDF2 — fallback quando pdftotext não está disponível.
    
    Instalar: pip3 install PyPDF2
    """
    try:
        import PyPDF2
        text_parts = []
        with open(pdf_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                try:
                    text_parts.append(page.extract_text() or "")
                except Exception:
                    text_parts.append("")
        return "\n".join(text_parts)
    except ImportError:
        return None
    except Exception:
        return None


def extract_with_pdfminer(pdf_path: str) -> str | None:
    """Extrai texto com pdfminer.six — mais robusto para layouts complexos.
    
    Instalar: pip3 install pdfminer.six
    """
    try:
        from pdfminer.high_level import extract_text
        return extract_text(pdf_path)
    except ImportError:
        return None
    except Exception:
        return None


def extract_with_docling(pdf_path: str) -> str | None:
    """Extrai PDF com reconhecimento de layout usando Docling.
    
    Melhor opção para PDFs técnicos com tabelas e fórmulas.
    Converte tabelas para markdown preservando a estrutura.
    Desvantagem: pacote pesado (~500MB), leva ~1,5s por página.
    Instalar: pip3 install docling
    """
    try:
        from docling.document_converter import DocumentConverter
        from docling.datamodel.pipeline_options import PdfPipelineOptions
        from docling.datamodel.base_models import InputFormat
        from docling.document_converter import PdfFormatOption

        pipeline_options = PdfPipelineOptions()
        pipeline_options.do_ocr = False            # Desativa OCR (mais rápido)
        pipeline_options.do_table_structure = True  # Ativa detecção de tabelas

        converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
            }
        )
        result = converter.convert(pdf_path)
        return result.document.export_to_markdown()
    except ImportError:
        return None
    except Exception:
        return None


# ─── Extratores EPUB ─────────────────────────────────────────────────────────

def extract_with_ebooklib(epub_path: str) -> str | None:
    """Extrai EPUB com ebooklib + BeautifulSoup4 — respeita ordem de leitura."""
    try:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup

        book = epub.read_epub(epub_path)
        parts = []
        # Itera apenas documentos (capítulos), não imagens ou CSS
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), "html.parser")
            parts.append(soup.get_text(separator="\n"))
        return "\n\n".join(parts)
    except ImportError:
        return None
    except Exception:
        return None


class _HTMLTextExtractor(html.parser.HTMLParser):
    """Conversor minimalista HTML → texto usando apenas a stdlib.
    
    Usado como fallback quando BeautifulSoup4 não está instalado.
    """

    SKIP_TAGS = {"script", "style", "head"}

    def __init__(self):
        super().__init__()
        self._parts: list[str] = []
        self._skip_depth = 0
        self._current_skip: str | None = None

    def handle_starttag(self, tag, attrs):
        if tag in self.SKIP_TAGS:
            self._skip_depth += 1
        # Adiciona quebra de linha antes de elementos de bloco
        if tag in ("p", "br", "h1", "h2", "h3", "h4", "h5", "h6", "li", "div"):
            self._parts.append("\n")

    def handle_endtag(self, tag):
        if tag in self.SKIP_TAGS and self._skip_depth:
            self._skip_depth -= 1

    def handle_data(self, data):
        if not self._skip_depth:
            self._parts.append(data)

    def get_text(self) -> str:
        return html.unescape("".join(self._parts))


def extract_with_zipfile(epub_path: str) -> str | None:
    """Extrai EPUB sem dependências — lê HTML dos arquivos ZIP diretamente."""
    try:
        with zipfile.ZipFile(epub_path) as zf:
            names = zf.namelist()
            # Lê o OPF para obter a ordem de leitura (spine)
            spine_order: list[str] = []
            opf_files = [n for n in names if n.endswith(".opf")]
            if opf_files:
                opf_text = zf.read(opf_files[0]).decode("utf-8", errors="replace")
                spine_order = re.findall(r'href=["\']([^"\']+\.(?:xhtml|html))["\']', opf_text)

            html_files = spine_order or sorted(
                n for n in names if n.endswith((".html", ".xhtml"))
            )
            if not html_files:
                return None

            parts = []
            for name in html_files:
                try:
                    raw = zf.read(name).decode("utf-8", errors="replace")
                    parser = _HTMLTextExtractor()
                    parser.feed(raw)
                    parts.append(parser.get_text())
                except Exception:
                    continue
            return "\n\n".join(parts) if parts else None
    except Exception:
        return None


def extract_epub(epub_path: str) -> tuple[str, str]:
    """Extrai EPUB tentando métodos em ordem de qualidade."""
    print("Trying ebooklib + BeautifulSoup4...", end=" ", flush=True)
    text = extract_with_ebooklib(epub_path)
    if text and text.strip():
        print("OK")
        return text, "ebooklib"
    print("not available")
    print("Trying stdlib zipfile parser...", end=" ", flush=True)
    text = extract_with_zipfile(epub_path)
    if text and text.strip():
        print("OK")
        return text, "zipfile"
    print("FAILED")
    print(
        "\nERROR: Could not extract text from EPUB.\n"
        "Install ebooklib + beautifulsoup4 for best results:\n"
        "  pip3 install ebooklib beautifulsoup4",
        file=sys.stderr,
    )
    sys.exit(1)


# ─── Contadores ──────────────────────────────────────────────────────────────

def count_epub_chapters(epub_path: str) -> int:
    """Conta itens no spine do EPUB — aproximação do número de capítulos."""
    try:
        with zipfile.ZipFile(epub_path) as zf:
            opf_files = [n for n in zf.namelist() if n.endswith(".opf")]
            if not opf_files:
                return 0
            opf_text = zf.read(opf_files[0]).decode("utf-8", errors="replace")
            return len(re.findall(r'<itemref\b', opf_text))
    except Exception:
        return 0


def count_pages(pdf_path: str) -> int:
    """Conta páginas de PDF — tenta pdfinfo primeiro, depois PyPDF2."""
    if shutil.which("pdfinfo"):
        try:
            result = subprocess.run(
                ["pdfinfo", pdf_path], capture_output=True, text=True, timeout=15
            )
            for line in result.stdout.splitlines():
                if line.startswith("Pages:"):
                    return int(line.split(":")[1].strip())
        except Exception:
            pass
    try:
        import PyPDF2
        with open(pdf_path, "rb") as f:
            return len(PyPDF2.PdfReader(f).pages)
    except Exception:
        return 0


# ─── Detecção de estrutura ───────────────────────────────────────────────────

def detect_structure(text: str) -> dict:
    """Detecta capítulos e sumário no documento.

    Mantido igual ao original do book-to-skill — simples e robusto.
    O Claude Code (SKILL.md) faz a análise estrutural real do conteúdo.
    Este script só fornece dicas básicas para o SKILL.md orientar o Claude.
    """
    lines = text[:50000].splitlines()

    # Detecta headings de capítulos no formato livro
    # Ex: "Chapter 1", "1. Introduction", "CHAPTER 2"
    chapter_pattern = re.compile(
        r"^\s*(chapter\s+\d+|CHAPTER\s+\d+|ch\.\s*\d+|\d+\.\s+[A-Z])",
        re.IGNORECASE
    )
    chapters_found = [l.strip() for l in lines if chapter_pattern.match(l)]

    # Detecta sumário (Table of Contents)
    # Exige que a palavra apareça sozinha na linha para evitar falsos positivos
    toc_pattern = re.compile(
        r"^\s*(?:table of contents|contents|índice|sumário)\s*$",
        re.IGNORECASE | re.MULTILINE,
    )
    has_toc = bool(toc_pattern.search(text[:30000]))

    return {
        "chapters_detected": len(chapters_found),
        "chapter_headings_sample": chapters_found[:10],
        "has_toc": has_toc,
    }


# ─── Função principal ────────────────────────────────────────────────────────

def main():
    """Ponto de entrada do script.

    Uso: extract.py <caminho> [--mode technical|text] [--install-missing ask|yes|no]

    Exemplos:
      python3 extract.py artigo.pdf
      python3 extract.py artigo.pdf --mode technical
      python3 extract.py artigo.pdf --install-missing yes
    """
    if len(sys.argv) < 2:
        print("Uso: extract.py <caminho-do-documento> [--mode technical|text] [--install-missing ask|yes|no]", file=sys.stderr)
        print(f"Formatos suportados: {supported_formats_message()}", file=sys.stderr)
        sys.exit(1)

    input_path = sys.argv[1]
    install_mode = normalize_install_mode(sys.argv)

    # Lê o modo de extração (technical = Docling, text = pdftotext/PyPDF2/pdfminer)
    extraction_mode = "text"
    if "--mode" in sys.argv:
        idx = sys.argv.index("--mode")
        if idx + 1 < len(sys.argv):
            extraction_mode = sys.argv[idx + 1].lower()
    if extraction_mode not in ("technical", "text"):
        extraction_mode = "text"

    if not os.path.exists(input_path):
        print(f"ERRO: Arquivo não encontrado: {input_path}", file=sys.stderr)
        sys.exit(1)

    input_file = Path(input_path)
    ext = input_file.suffix.lower()
    document_format = ext.lstrip(".")

    # Detecção por magic bytes para arquivos sem extensão reconhecida
    # PDF começa com "%PDF", EPUB e DOCX são ZIPs (começam com "PK")
    if ext not in SUPPORTED_EXTENSIONS:
        with open(input_path, "rb") as f:
            header = f.read(8)
        if header[:4] == b"%PDF":
            ext = ".pdf"
            document_format = "pdf"
        elif header[:2] == b"PK":
            try:
                with zipfile.ZipFile(input_path) as zf:
                    names = set(zf.namelist())
                    if "mimetype" in names and zf.read("mimetype").startswith(b"application/epub"):
                        ext = ".epub"
                        document_format = "epub"
                    elif "word/document.xml" in names:
                        ext = ".docx"
                        document_format = "docx"
                    else:
                        print(
                            f"ERRO: Formato ZIP não suportado '{input_file.name}'. Suportados: {supported_formats_message()}",
                            file=sys.stderr,
                        )
                        sys.exit(1)
            except (zipfile.BadZipFile, KeyError, OSError):
                print(
                    f"ERRO: ZIP inválido '{input_file.name}'. Suportados: {supported_formats_message()}",
                    file=sys.stderr,
                )
                sys.exit(1)
        else:
            print(
                f"ERRO: Formato não suportado '{ext or '<nenhum>'}'. Suportados: {supported_formats_message()}",
                file=sys.stderr,
            )
            sys.exit(1)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    prepare_dependencies(ext, extraction_mode, install_mode)

    # MOBI/AZW requerem Calibre instalado no sistema (não via pip)
    if ext in CALIBRE_EBOOK_EXTENSIONS and not shutil.which("ebook-convert"):
        print(
            "ERRO: MOBI/AZW/AZW3 requer o Calibre. "
            "Instale em https://calibre-ebook.com/download e garanta que "
            "ebook-convert esteja no PATH.",
            file=sys.stderr,
        )
        sys.exit(1)

    # ── Roteamento por formato ──
    if ext == ".epub":
        print(f"Extraindo EPUB: {input_path}")
        text, method = extract_epub(input_path)
        pages = count_epub_chapters(input_path)
        pages_label = "spine_items"

    elif ext == ".pdf":
        print(f"Extraindo PDF: {input_path}")
        if extraction_mode == "technical":
            # Modo técnico: Docling preserva tabelas e fórmulas como markdown
            print("Modo: técnico — usando Docling (reconhecimento de layout)...", end=" ", flush=True)
            text = extract_with_docling(input_path)
            if text:
                method = "docling"
                print("OK")
            else:
                # Docling não disponível — cai para a cadeia de texto
                print("não disponível, usando pdftotext como fallback")
                extraction_mode = "text"

        if extraction_mode == "text":
            # Modo texto: cadeia pdftotext → PyPDF2 → pdfminer
            print("Modo: texto — usando pdftotext...")
            print("Tentando pdftotext...", end=" ", flush=True)
            text = extract_with_pdftotext(input_path)
            if text:
                method = "pdftotext"
                print("OK")
            else:
                print("não disponível")
                print("Tentando PyPDF2...", end=" ", flush=True)
                text = extract_with_pypdf2(input_path)
                if text:
                    method = "PyPDF2"
                    print("OK")
                else:
                    print("não disponível")
                    print("Tentando pdfminer.six...", end=" ", flush=True)
                    text = extract_with_pdfminer(input_path)
                    if text:
                        method = "pdfminer"
                        print("OK")
                    else:
                        print("FALHOU")
                        print(
                            "\nERRO: Não foi possível extrair texto do PDF.\n"
                            "Instale uma das opções abaixo:\n"
                            "  sudo apt install poppler-utils  (recomendado)\n"
                            "  pip3 install PyPDF2\n"
                            "  pip3 install pdfminer.six",
                            file=sys.stderr,
                        )
                        sys.exit(1)
        pages = count_pages(input_path)
        pages_label = "pages"

    elif ext in TEXT_EXTENSIONS:
        print(f"Extraindo documento de texto: {input_path}")
        text = read_text_file(input_path)
        if text is None or not text.strip():
            print("ERRO: Não foi possível ler o documento de texto", file=sys.stderr)
            sys.exit(1)
        method = "plain-text"
        pages = 0
        pages_label = "sections"

    elif ext in HTML_EXTENSIONS:
        print(f"Extraindo HTML: {input_path}")
        text = extract_html_file(input_path)
        if text is None or not text.strip():
            print("ERRO: Não foi possível extrair texto do HTML", file=sys.stderr)
            sys.exit(1)
        method = "html-parser"
        pages = 0
        pages_label = "sections"

    elif ext == ".docx":
        print(f"Extraindo DOCX: {input_path}")
        text, method = extract_docx(input_path)
        pages = 0
        pages_label = "sections"

    elif ext == ".rtf":
        print(f"Extraindo RTF: {input_path}")
        text, method = extract_rtf(input_path)
        pages = 0
        pages_label = "sections"

    elif ext in CALIBRE_EBOOK_EXTENSIONS:
        print(f"Extraindo ebook com Calibre: {input_path}")
        text = extract_with_ebook_convert(input_path)
        if text is None or not text.strip():
            print(
                f"ERRO: Não foi possível extrair texto de {ext}. "
                "Instale o Calibre e garanta que ebook-convert esteja no PATH.",
                file=sys.stderr,
            )
            sys.exit(1)
        method = "ebook-convert"
        pages = 0
        pages_label = "sections"

    else:
        print(
            f"ERRO: Formato não suportado '{ext}'. Suportados: {supported_formats_message()}",
            file=sys.stderr,
        )
        sys.exit(1)

    # ── Salva os resultados ──
    OUTPUT_TEXT.write_text(text, encoding="utf-8")

    tokens = estimate_tokens(text)
    structure = detect_structure(text)
    file_size_mb = os.path.getsize(input_path) / (1024 * 1024)

    # metadata.json é lido pelo SKILL.md para:
    # - estimated_tokens → estimativa de custo
    # - chapters_detected / has_toc → orientar o Claude na análise de estrutura
    metadata = {
        "source_file": str(Path(input_path).resolve()),
        "filename": Path(input_path).name,
        "format": document_format,
        "extraction_method": method,
        "extraction_mode": extraction_mode,
        "file_size_mb": round(file_size_mb, 2),
        pages_label: pages,
        "chars": len(text),
        "words": len(text.split()),
        "estimated_tokens": tokens,
        "estimated_tokens_human": f"~{tokens // 1000}K",
        "output_text": str(OUTPUT_TEXT),
        **structure,
    }

    OUTPUT_META.write_text(json.dumps(metadata, indent=2, ensure_ascii=False))

    # ── Relatório final no terminal ──
    page_label = {
        "spine_items": "Itens no spine",
        "pages": "Páginas",
        "sections": "Seções",
    }.get(pages_label, pages_label.replace("_", " ").title())

    print("\nExtração concluída:")
    print(f"   Formato   : {document_format.upper()}")
    print(f"   Método    : {method}")
    print(f"   {page_label:12}: {pages}")
    print(f"   Palavras  : {len(text.split()):,}")
    print(f"   Tokens ~  : ~{tokens // 1000}K")
    print(f"   Capítulos : {structure['chapters_detected']} detectados")
    print(f"   Sumário   : {'sim' if structure['has_toc'] else 'não detectado'}")
    if not structure["has_toc"]:
        print(
            "   AVISO     : Sem sumário — o Claude mapeará seções por varredura de headings."
        )
    print(f"\n   Texto     → {OUTPUT_TEXT}")
    print(f"   Metadados → {OUTPUT_META}")


if __name__ == "__main__":
    main()